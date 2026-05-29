from __future__ import annotations

from datetime import date
from pathlib import Path
from shutil import copy2

from docx import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import HTTPException, status
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.acknowledgement import Acknowledgement
from app.models.approval_step import ApprovalStep
from app.models.document import Document
from app.models.document_file import DocumentFile
from app.models.enums import AcknowledgementStatus, ApprovalStatus, DocumentFileKind, DocumentType
from app.models.user import User
from app.services.markdown_docx import markdown_docx_renderer
from app.services.pdf import pdf_conversion_service
from app.services.storage import ensure_parent, sha256sum, storage_root


class DocumentGenerationService:
    async def regenerate_document_files(self, db: AsyncSession, document: Document, current_user: User) -> list[DocumentFile]:
        if document.type == DocumentType.ORDER:
            return await self.generate_order_docx(db, document, current_user)
        if document.type == DocumentType.INSTRUCTION:
            return await self.generate_instruction_docx(db, document, current_user)
        if document.type == DocumentType.ORDER_EXTRACT:
            return await self.generate_order_extract_docx(db, document, current_user)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Generation is not supported for this document type")

    async def generate_order_docx(self, db: AsyncSession, document: Document, current_user: User) -> list[DocumentFile]:
        return await self._generate_main_document(db, document, current_user, title_text="П Р И К А З", command_text="П Р И К А З Ы В А Ю:")

    async def generate_instruction_docx(self, db: AsyncSession, document: Document, current_user: User) -> list[DocumentFile]:
        return await self._generate_main_document(
            db,
            document,
            current_user,
            title_text="П Р И К А З А Н И Е",
            command_text="П Р И К А З Ы В А Ю:",
        )

    async def generate_order_extract_docx(self, db: AsyncSession, document: Document, current_user: User) -> list[DocumentFile]:
        await self._delete_existing_generated_files(db, document.id, {DocumentFileKind.EXTRACT_DOCX, DocumentFileKind.EXTRACT_PDF})

        doc = DocxDocument()
        self._apply_base_style(doc)
        data = document.structured_data
        self._center_paragraph(doc, document.organization_name, bold=True)
        self._center_paragraph(doc, "В Ы П И С К А И З П Р И К А З А", bold=True, size=16)
        self._center_paragraph(doc, document.signer_position, bold=True)
        self._center_paragraph(doc, f"от {self._format_date(document.document_date or date.today())} {document.city}")
        markdown_docx_renderer.render(doc, data.get("certifier_position", ""))
        self._body_paragraph(doc, document.title, bold=True)
        for index, item in enumerate(data.get("extracted_items", []), start=1):
            markdown_docx_renderer.render_numbered_item(doc, index, item)
        self._body_paragraph(doc, "")
        markdown_docx_renderer.render(doc, f"Выписка верна: {data.get('certifier_position', '')}")
        markdown_docx_renderer.render(doc, data.get("certifier_name", ""))

        docx_path, pdf_path = self._document_paths(document)
        ensure_parent(docx_path)
        doc.save(docx_path)
        converted_pdf = pdf_conversion_service.convert_docx_to_pdf(docx_path, pdf_path.parent)
        if converted_pdf != pdf_path:
            copy2(converted_pdf, pdf_path)
            converted_pdf.unlink(missing_ok=True)

        return await self._persist_generated_files(
            db,
            document,
            current_user,
            [
                (docx_path, DocumentFileKind.EXTRACT_DOCX, True, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                (pdf_path, DocumentFileKind.EXTRACT_PDF, True, "application/pdf"),
            ],
        )

    async def _generate_main_document(
        self,
        db: AsyncSession,
        document: Document,
        current_user: User,
        *,
        title_text: str,
        command_text: str,
    ) -> list[DocumentFile]:
        document = await self._load_document_for_generation(db, document.id)
        await self._delete_existing_generated_files(db, document.id, {DocumentFileKind.GENERATED_DOCX, DocumentFileKind.GENERATED_PDF})
        doc = DocxDocument()
        self._apply_base_style(doc)
        data = document.structured_data

        self._center_paragraph(doc, document.organization_name, bold=True)
        self._center_paragraph(doc, title_text, bold=True, size=16)
        number_line = document.registered_number or "без номера"
        date_line = self._format_date(document.document_date or date.today())
        self._center_paragraph(doc, f"№ {number_line} от {date_line}", bold=True)
        self._center_paragraph(doc, document.city)
        subject = data.get("order_subject") or data.get("instruction_subject") or document.title
        self._center_paragraph(doc, subject, bold=True)
        if data.get("legal_basis_text"):
            markdown_docx_renderer.render(doc, data["legal_basis_text"])
        if data.get("purpose_text"):
            markdown_docx_renderer.render(doc, data["purpose_text"])
        self._body_paragraph(doc, command_text, bold=True)
        for index, item in enumerate(data.get("order_items", data.get("instruction_items", [])), start=1):
            markdown_docx_renderer.render_numbered_item(doc, index, item)
        markdown_docx_renderer.render(doc, f"Контроль исполнения возложить на {data.get('control_assignee_text', '')}.")
        self._body_paragraph(doc, "")
        self._body_paragraph(doc, document.signer_position)
        self._body_paragraph(doc, document.signer_name)

        if document.type == DocumentType.ORDER and (document.executor_name or data.get("executor_name")):
            self._body_paragraph(doc, "")
            self._body_paragraph(doc, f"Исп.: {document.executor_name or data.get('executor_name')}")
        if document.type == DocumentType.ORDER and (document.executor_phone or data.get("executor_phone")):
            self._body_paragraph(doc, f"тел. {document.executor_phone or data.get('executor_phone')}")

        if document.type == DocumentType.ORDER:
            preview_users = await self._load_preview_users(db, data)
            self._append_approval_sheet(doc, document, preview_users)
            self._append_acknowledgement_sheet(doc, document, preview_users)

        docx_path, pdf_path = self._document_paths(document)
        ensure_parent(docx_path)
        doc.save(docx_path)
        converted_pdf = pdf_conversion_service.convert_docx_to_pdf(docx_path, pdf_path.parent)
        if converted_pdf != pdf_path:
            copy2(converted_pdf, pdf_path)
            converted_pdf.unlink(missing_ok=True)

        return await self._persist_generated_files(
            db,
            document,
            current_user,
            [
                (docx_path, DocumentFileKind.GENERATED_DOCX, False, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                (pdf_path, DocumentFileKind.GENERATED_PDF, False, "application/pdf"),
            ],
        )

    async def _delete_existing_generated_files(
        self,
        db: AsyncSession,
        document_id: int,
        kinds: set[DocumentFileKind],
    ) -> None:
        result = await db.execute(select(DocumentFile).where(DocumentFile.document_id == document_id, DocumentFile.kind.in_(tuple(kinds))))
        for file_entry in result.scalars().all():
            Path(file_entry.storage_path).unlink(missing_ok=True)
        await db.execute(delete(DocumentFile).where(DocumentFile.document_id == document_id, DocumentFile.kind.in_(tuple(kinds))))
        await db.flush()

    async def _persist_generated_files(
        self,
        db: AsyncSession,
        document: Document,
        current_user: User,
        files: list[tuple[Path, DocumentFileKind, bool, str]],
    ) -> list[DocumentFile]:
        stored_files: list[DocumentFile] = []
        for path, kind, download_allowed, mime_type in files:
            file_entry = DocumentFile(
                document_id=document.id,
                version=document.current_version,
                original_filename=path.name,
                storage_path=str(path),
                mime_type=mime_type,
                size_bytes=path.stat().st_size,
                sha256=sha256sum(path),
                kind=kind,
                is_download_allowed=download_allowed,
                created_by=current_user.id,
            )
            db.add(file_entry)
            stored_files.append(file_entry)
        await db.flush()
        return stored_files

    def _document_paths(self, document: Document) -> tuple[Path, Path]:
        base = storage_root()
        doc_date = document.document_date or date.today()
        if document.type == DocumentType.ORDER_EXTRACT:
            directory = base / "extracts" / f"{doc_date.year}" / f"{doc_date.month:02d}" / str(document.id)
        else:
            directory = (
                base
                / "documents"
                / f"{doc_date.year}"
                / f"{doc_date.month:02d}"
                / str(document.id)
                / str(document.current_version)
            )
        return directory / "generated.docx", directory / "generated.pdf"

    def _apply_base_style(self, document: DocxDocument) -> None:
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)

    def _center_paragraph(self, document: DocxDocument, text: str, *, bold: bool = False, size: int = 14) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)

    def _body_paragraph(self, document: DocxDocument, text: str, *, bold: bool = False) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(20)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    def _format_date(self, value: date) -> str:
        return value.strftime("%d.%m.%Y")

    def _append_acknowledgement_sheet(self, document: DocxDocument, source: Document, preview_users: dict[int, User]) -> None:
        self._sheet_title(document, "ЛИСТ ОЗНАКОМЛЕНИЯ")
        table = document.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["№", "ФИО", "Должность", "Статус", "Дата"]
        self._fill_header_row(table.rows[0].cells, headers)
        acknowledgements = source.acknowledgements
        if not acknowledgements:
            fallback_ids = source.structured_data.get("acknowledgement_people", [])
            acknowledgements = [
                type(
                    "AckPreview",
                    (),
                    {
                        "user_id": user_id,
                        "status": AcknowledgementStatus.PENDING,
                        "acknowledged_at": None,
                        "user": preview_users.get(user_id),
                    },
                )()
                for user_id in fallback_ids
            ]
        for index, ack in enumerate(acknowledgements, start=1):
            row = table.add_row().cells
            user = ack.user or preview_users.get(ack.user_id)
            row[0].text = str(index)
            row[1].text = user.full_name if user else f"Пользователь #{ack.user_id}"
            row[2].text = user.position if user and user.position else "-"
            row[3].text = ack.status.value
            row[4].text = ack.acknowledged_at.strftime("%d.%m.%Y %H:%M") if ack.acknowledged_at else ""
            self._format_row(row)

    def _append_approval_sheet(self, document: DocxDocument, source: Document, preview_users: dict[int, User]) -> None:
        self._sheet_title(document, "ЛИСТ СОГЛАСОВАНИЯ")
        table = document.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["№", "Должность, ФИО", "Статус", "Комментарий", "Дата"]
        self._fill_header_row(table.rows[0].cells, headers)
        approval_steps = source.approval_steps
        if not approval_steps:
            fallback_ids = source.structured_data.get("approval_people", [])
            approval_steps = [
                type(
                    "StepPreview",
                    (),
                    {
                        "step_order": index,
                        "approver_id": approver_id,
                        "status": ApprovalStatus.WAITING if index == 1 else ApprovalStatus.PENDING,
                        "comment": "",
                        "acted_at": None,
                        "approver": preview_users.get(approver_id),
                    },
                )()
                for index, approver_id in enumerate(fallback_ids, start=1)
            ]
        for step in approval_steps:
            row = table.add_row().cells
            approver = step.approver or preview_users.get(step.approver_id)
            row[0].text = str(step.step_order)
            row[1].text = self._person_label(approver, step.approver_id)
            row[2].text = step.status.value
            row[3].text = step.comment or ""
            row[4].text = step.acted_at.strftime("%d.%m.%Y %H:%M") if step.acted_at else ""
            self._format_row(row)

    async def _load_preview_users(self, db: AsyncSession, structured_data: dict) -> dict[int, User]:
        user_ids = list(
            dict.fromkeys(
                [*structured_data.get("approval_people", []), *structured_data.get("acknowledgement_people", [])]
            )
        )
        if not user_ids:
            return {}
        result = await db.execute(select(User).where(User.id.in_(user_ids)))
        users = result.scalars().all()
        return {user.id: user for user in users}

    async def _load_document_for_generation(self, db: AsyncSession, document_id: int) -> Document:
        result = await db.execute(
            select(Document)
            .options(
                selectinload(Document.approval_steps).selectinload(ApprovalStep.approver),
                selectinload(Document.acknowledgements).selectinload(Acknowledgement.user),
            )
            .where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        if document is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
        return document

    def _sheet_title(self, document: DocxDocument, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.page_break_before = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    def _fill_header_row(self, cells, values: list[str]) -> None:
        for cell, value in zip(cells, values, strict=False):
            cell.text = value
            self._format_cell(cell, bold=True, centered=True)

    def _format_row(self, cells) -> None:
        for index, cell in enumerate(cells):
            self._format_cell(cell, centered=index != 1)

    def _format_cell(self, cell, *, bold: bool = False, centered: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = bold
        if not paragraph.runs:
            run = paragraph.add_run("")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = bold

    def _person_label(self, approver: User | None, approver_id: int) -> str:
        if approver is None:
            return f"Пользователь #{approver_id}"
        position = f"{approver.position}, " if approver.position else ""
        return f"{position}{approver.full_name}"


document_generation_service = DocumentGenerationService()
