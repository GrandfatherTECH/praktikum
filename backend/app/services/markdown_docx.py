from __future__ import annotations

from dataclasses import dataclass

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt
from markdown_it.token import Token


@dataclass(slots=True)
class RenderContext:
    prefix: str | None = None
    nested: bool = False


class MarkdownDocxRenderer:
    def __init__(self) -> None:
        self._md = MarkdownIt("commonmark", {"html": False, "linkify": False, "typographer": False})

    def render(self, document: DocxDocument, markdown: str, *, first_line_indent: float = 20) -> None:
        tokens = self._md.parse(markdown or "")
        if not tokens:
            return
        self._render_blocks(document, tokens, 0, len(tokens), first_line_indent=first_line_indent, context=RenderContext())

    def render_numbered_item(self, document: DocxDocument, index: int, markdown: str) -> None:
        tokens = self._md.parse(markdown or "")
        if not tokens:
            self._render_paragraph(document, [], first_line_indent=20, prefix=f"{index}. ")
            return
        self._render_blocks(document, tokens, 0, len(tokens), first_line_indent=20, context=RenderContext(prefix=f"{index}. "))

    def _render_blocks(
        self,
        document: DocxDocument,
        tokens: list[Token],
        start: int,
        end: int,
        *,
        first_line_indent: float,
        context: RenderContext,
    ) -> int:
        index = start
        first_block = True
        while index < end:
            token = tokens[index]
            if token.type in {"paragraph_open", "heading_open"}:
                inline_token = tokens[index + 1]
                is_heading = token.type == "heading_open"
                prefix = context.prefix if first_block else None
                self._render_paragraph(
                    document,
                    inline_token.children or [],
                    first_line_indent=first_line_indent if not context.nested else 0,
                    left_indent=20 if context.nested else 0,
                    prefix=prefix,
                    bold=is_heading,
                    italic=False,
                    size=16 - int(token.tag[1]) if is_heading else 14,
                )
                index += 3
                first_block = False
                continue

            if token.type in {"bullet_list_open", "ordered_list_open", "blockquote_open"}:
                close_type = token.type.replace("_open", "_close")
                inner_end = self._find_matching_close(tokens, index, close_type)
                prefix = context.prefix if first_block else None
                if token.type == "blockquote_open":
                    self._render_blockquote(document, tokens, index + 1, inner_end, first_line_indent, prefix=prefix)
                else:
                    ordered = token.type == "ordered_list_open"
                    start_number = int(token.attrGet("start") or 1) if ordered else 1
                    self._render_list(
                        document,
                        tokens,
                        index + 1,
                        inner_end,
                        first_line_indent=first_line_indent,
                        ordered=ordered,
                        start_number=start_number,
                        prefix=prefix,
                        nested=context.nested,
                    )
                index = inner_end + 1
                first_block = False
                continue

            index += 1

        return index

    def _render_list(
        self,
        document: DocxDocument,
        tokens: list[Token],
        start: int,
        end: int,
        *,
        first_line_indent: float,
        ordered: bool,
        start_number: int,
        prefix: str | None,
        nested: bool,
    ) -> None:
        item_number = start_number
        index = start
        first_item = True
        while index < end:
            token = tokens[index]
            if token.type != "list_item_open":
                index += 1
                continue
            item_end = self._find_matching_close(tokens, index, "list_item_close")
            item_prefix = prefix if first_item else None
            marker = f"{item_number}. " if ordered else "• "
            if item_prefix:
                marker = f"{item_prefix}{marker}"
            self._render_list_item(
                document,
                tokens,
                index + 1,
                item_end,
                first_line_indent=first_line_indent,
                marker=marker,
                nested=nested,
            )
            if ordered:
                item_number += 1
            index = item_end + 1
            first_item = False

    def _render_list_item(
        self,
        document: DocxDocument,
        tokens: list[Token],
        start: int,
        end: int,
        *,
        first_line_indent: float,
        marker: str,
        nested: bool,
    ) -> None:
        index = start
        first_block = True
        while index < end:
            token = tokens[index]
            if token.type in {"paragraph_open", "heading_open"}:
                inline_token = tokens[index + 1]
                self._render_paragraph(
                    document,
                    inline_token.children or [],
                    first_line_indent=0 if first_block else 0,
                    left_indent=20 if not nested else 36,
                    prefix=marker if first_block else None,
                    bold=token.type == "heading_open",
                    size=16 - int(token.tag[1]) if token.type == "heading_open" else 14,
                )
                index += 3
                first_block = False
                continue

            if token.type in {"bullet_list_open", "ordered_list_open", "blockquote_open"}:
                close_type = token.type.replace("_open", "_close")
                inner_end = self._find_matching_close(tokens, index, close_type)
                if token.type == "blockquote_open":
                    self._render_blockquote(
                        document,
                        tokens,
                        index + 1,
                        inner_end,
                        first_line_indent,
                        prefix=marker if first_block else None,
                        left_indent=36 if not nested else 48,
                    )
                else:
                    ordered = token.type == "ordered_list_open"
                    start_number = int(token.attrGet("start") or 1) if ordered else 1
                    self._render_list(
                        document,
                        tokens,
                        index + 1,
                        inner_end,
                        first_line_indent=first_line_indent,
                        ordered=ordered,
                        start_number=start_number,
                        prefix=marker if first_block else None,
                        nested=True,
                    )
                index = inner_end + 1
                first_block = False
                continue

            index += 1

    def _render_blockquote(
        self,
        document: DocxDocument,
        tokens: list[Token],
        start: int,
        end: int,
        first_line_indent: float,
        *,
        prefix: str | None,
        left_indent: float = 24,
    ) -> None:
        index = start
        first_block = True
        while index < end:
            token = tokens[index]
            if token.type in {"paragraph_open", "heading_open"}:
                inline_token = tokens[index + 1]
                self._render_paragraph(
                    document,
                    inline_token.children or [],
                    first_line_indent=0,
                    left_indent=left_indent,
                    prefix=prefix if first_block else None,
                    bold=token.type == "heading_open",
                    italic=True,
                    size=16 - int(token.tag[1]) if token.type == "heading_open" else 14,
                )
                index += 3
                first_block = False
                continue
            index += 1

    def _render_paragraph(
        self,
        document: DocxDocument,
        inline_tokens: list[Token],
        *,
        first_line_indent: float,
        left_indent: float = 0,
        prefix: str | None = None,
        bold: bool = False,
        italic: bool = False,
        size: int = 14,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
        paragraph.paragraph_format.left_indent = Pt(left_indent)
        paragraph.paragraph_format.space_after = Pt(0)
        if prefix:
            self._append_run(paragraph, prefix, bold=bold, italic=italic, size=size)
        self._append_inline_tokens(paragraph, inline_tokens, bold=bold, italic=italic, size=size)

    def _append_inline_tokens(
        self,
        paragraph: Paragraph,
        inline_tokens: list[Token],
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 14,
    ) -> None:
        bold_stack = 0
        italic_stack = 0
        code_stack = 0

        for token in inline_tokens:
            if token.type == "strong_open":
                bold_stack += 1
                continue
            if token.type == "strong_close":
                bold_stack = max(0, bold_stack - 1)
                continue
            if token.type == "em_open":
                italic_stack += 1
                continue
            if token.type == "em_close":
                italic_stack = max(0, italic_stack - 1)
                continue
            if token.type == "code_inline":
                self._append_run(
                    paragraph,
                    token.content,
                    bold=bold or bold_stack > 0,
                    italic=italic or italic_stack > 0,
                    size=size,
                    font_name="Courier New",
                )
                continue
            if token.type == "text":
                self._append_run(
                    paragraph,
                    token.content,
                    bold=bold or bold_stack > 0,
                    italic=italic or italic_stack > 0,
                    size=size,
                )
                continue
            if token.type in {"softbreak", "hardbreak"}:
                paragraph.add_run("\n")
                continue
            if token.type == "html_inline":
                continue
            if token.type.endswith("_open"):
                code_stack += 1
                continue
            if token.type.endswith("_close"):
                code_stack = max(0, code_stack - 1)
                continue
            if token.content:
                self._append_run(
                    paragraph,
                    token.content,
                    bold=bold or bold_stack > 0,
                    italic=italic or italic_stack > 0,
                    size=size,
                    font_name="Courier New" if code_stack > 0 else "Times New Roman",
                )

    def _append_run(
        self,
        paragraph: Paragraph,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 14,
        font_name: str = "Times New Roman",
    ) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(size)

    def _find_matching_close(self, tokens: list[Token], start: int, close_type: str) -> int:
        depth = 0
        open_type = tokens[start].type
        for index in range(start, len(tokens)):
            token = tokens[index]
            if token.type == open_type:
                depth += 1
            elif token.type == close_type:
                depth -= 1
                if depth == 0:
                    return index
        raise ValueError(f"Unbalanced markdown token stream for {open_type}")


markdown_docx_renderer = MarkdownDocxRenderer()
