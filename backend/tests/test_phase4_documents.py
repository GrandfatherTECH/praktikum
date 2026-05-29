from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.models.audit_log import AuditLog
from app.models.document import Document
from app.models.document_file import DocumentFile
from app.models.enums import ApprovalStatus, DocumentFileKind, DocumentStatus
from app.services.pdf import pdf_conversion_service


PDF_BYTES = b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF\n"


@pytest.fixture(autouse=True)
def fake_pdf_conversion(monkeypatch):
    def _convert(source_docx: Path, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = output_dir / f"{source_docx.stem}.pdf"
        pdf_path.write_bytes(PDF_BYTES)
        return pdf_path

    monkeypatch.setattr(pdf_conversion_service, "convert_docx_to_pdf", _convert)


async def do_login(client, username: str, password: str) -> None:
    response = await client.post("/api/v1/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200


def build_order_payload(users_fixture) -> dict:
    return {
        "title": "О назначении ответственных",
        "department_id": users_fixture["department"].id,
        "document_date": "2026-05-28",
        "city": "г. Екатеринбург",
        "organization_name": "ФГКОУ «Екатеринбургское суворовское военное училище»",
        "signer_position": "Начальник училища",
        "signer_name": "А. Кукарцев",
        "executor_name": "Ю. Ятченко",
        "executor_phone": "2-74",
        "structured_data": {
            "order_subject": "О назначении ответственных",
            "legal_basis_text": "**В целях** организации работы.\n\n> Основание довести до исполнителей.",
            "purpose_text": "Для обеспечения *контроля* исполнения.",
            "order_items": [
                "Назначить **ответственного** за архив.\n- Подготовить ключи\n- Сверить журнал",
                "Обеспечить еженедельный контроль.\n1. Проверять отчеты\n2. Докладывать начальнику",
            ],
            "control_assignee_text": "заместителя начальника училища",
            "approval_people": [users_fixture["chief"].id, users_fixture["dept_head"].id],
            "acknowledgement_people": [users_fixture["employee"].id],
            "acknowledgement_departments": [users_fixture["department"].id],
        },
    }


async def test_create_structured_order(client, users_fixture):
    await do_login(client, "admin", "admin12345")
    response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    assert response.status_code == 201
    data = response.json()
    assert data["type"] == "ORDER"
    assert data["structured_data"]["order_subject"] == "О назначении ответственных"
    assert data["acknowledgements"]


async def test_generate_order_files_and_preview_access(client, users_fixture, db_session):
    await do_login(client, "admin", "admin12345")
    create_response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    document_id = create_response.json()["id"]

    generate_response = await client.post(f"/api/v1/documents/orders/{document_id}/generate")
    assert generate_response.status_code == 200

    result = await db_session.execute(select(DocumentFile).where(DocumentFile.document_id == document_id))
    files = result.scalars().all()
    assert {item.kind for item in files} == {DocumentFileKind.GENERATED_DOCX, DocumentFileKind.GENERATED_PDF}
    assert all(item.is_download_allowed is False for item in files)

    preview_response = await client.get(f"/api/v1/documents/{document_id}/preview")
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"] == "application/pdf"

    client.cookies.clear()
    await do_login(client, "personnel", "personnel12345")
    forbidden_preview = await client.get(f"/api/v1/documents/{document_id}/preview")
    assert forbidden_preview.status_code == 403


async def test_send_for_approval_and_sequential_rule(client, users_fixture, db_session):
    await do_login(client, "admin", "admin12345")
    create_response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    document_id = create_response.json()["id"]
    send_response = await client.post(
        f"/api/v1/documents/orders/{document_id}/send-for-approval",
        json={"approver_ids": [users_fixture["chief"].id, users_fixture["dept_head"].id]},
    )
    assert send_response.status_code == 200
    assert send_response.json()["status"] == "ON_APPROVAL"

    client.cookies.clear()
    await do_login(client, "dept_head", "depthead12345")
    forbidden = await client.post(f"/api/v1/documents/{document_id}/approve", json={})
    assert forbidden.status_code == 403

    client.cookies.clear()
    await do_login(client, "chief", "chief12345")
    approved = await client.post(f"/api/v1/documents/{document_id}/approve", json={})
    assert approved.status_code == 200
    assert approved.json()["approval_steps"][0]["status"] == "APPROVED"
    assert approved.json()["approval_steps"][1]["status"] == "WAITING"

    result = await db_session.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one()
    assert document.status == DocumentStatus.ON_APPROVAL


async def test_return_for_revision_requires_comment(client, users_fixture):
    await do_login(client, "admin", "admin12345")
    create_response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    document_id = create_response.json()["id"]
    await client.post(
        f"/api/v1/documents/orders/{document_id}/send-for-approval",
        json={"approver_ids": [users_fixture["chief"].id]},
    )

    client.cookies.clear()
    await do_login(client, "chief", "chief12345")
    invalid = await client.post(f"/api/v1/documents/{document_id}/return-for-revision", json={"comment": ""})
    assert invalid.status_code == 422

    valid = await client.post(
        f"/api/v1/documents/{document_id}/return-for-revision",
        json={"comment": "Нужно уточнить исполнителя."},
    )
    assert valid.status_code == 200
    assert valid.json()["status"] == "REVISION_REQUIRED"


async def test_acknowledgement_and_generated_docx_contains_lists(client, users_fixture, db_session):
    await do_login(client, "admin", "admin12345")
    create_response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    document_id = create_response.json()["id"]
    await client.post(
        f"/api/v1/documents/orders/{document_id}/send-for-approval",
        json={"approver_ids": [users_fixture["chief"].id]},
    )

    client.cookies.clear()
    await do_login(client, "chief", "chief12345")
    await client.post(f"/api/v1/documents/{document_id}/approve", json={})

    client.cookies.clear()
    await do_login(client, "admin", "admin12345")
    send_ack = await client.post(
        f"/api/v1/documents/{document_id}/send-for-acknowledgement",
        json={"user_ids": [users_fixture["employee"].id], "department_ids": []},
    )
    assert send_ack.status_code == 200

    client.cookies.clear()
    await do_login(client, "employee", "employee12345")
    ack = await client.post(f"/api/v1/documents/{document_id}/acknowledge")
    assert ack.status_code == 200
    assert ack.json()["status"] == "ACKNOWLEDGEMENT_COMPLETED"

    result = await db_session.execute(
        select(DocumentFile).where(
            DocumentFile.document_id == document_id,
            DocumentFile.kind == DocumentFileKind.GENERATED_DOCX,
        )
    )
    docx_file = result.scalar_one()
    with zipfile.ZipFile(docx_file.storage_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "Лист согласования" in xml
    assert "Chief" in xml
    assert "Лист ознакомления" in xml
    assert "Employee" in xml

    doc = DocxDocument(docx_file.storage_path)
    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    assert any("• Подготовить ключи" in text for text in paragraph_texts)
    assert any("1. Проверять отчеты" in text for text in paragraph_texts)
    assert any("Основание довести до исполнителей." in text for text in paragraph_texts)

    assert any(run.text == "В целях" and run.bold for paragraph in doc.paragraphs for run in paragraph.runs)
    assert any(run.text == "контроля" and run.italic for paragraph in doc.paragraphs for run in paragraph.runs)


async def test_generate_extract_and_download_allowed(client, users_fixture):
    await do_login(client, "admin", "admin12345")
    create_response = await client.post("/api/v1/documents/orders", json=build_order_payload(users_fixture))
    document_id = create_response.json()["id"]
    await client.post(f"/api/v1/documents/orders/{document_id}/generate")
    extract_response = await client.post(
        f"/api/v1/documents/{document_id}/generate-extract",
        json={
            "extracted_items": ["Назначить ответственного за архив."],
            "certifier_position": "Зав. делопроизводством",
            "certifier_name": "Н. Иванкова",
            "extract_date": "2026-05-28",
        },
    )
    assert extract_response.status_code == 201
    extract = extract_response.json()
    assert extract["type"] == "ORDER_EXTRACT"
    downloadables = [item for item in extract["files"] if item["is_download_allowed"]]
    assert len(downloadables) == 2

    download = await client.get(f"/api/v1/documents/{extract['id']}/files/{downloadables[0]['id']}/download")
    assert download.status_code == 200


async def test_create_instruction(client, users_fixture):
    await do_login(client, "admin", "admin12345")
    response = await client.post(
        "/api/v1/documents/instructions",
        json={
            "title": "О порядке дежурства",
            "department_id": users_fixture["department"].id,
            "document_date": "2026-05-28",
            "city": "г. Екатеринбург",
            "organization_name": "ФГКОУ «Екатеринбургское суворовское военное училище»",
            "signer_position": "Начальник училища",
            "signer_name": "А. Кукарцев",
            "structured_data": {
                "instruction_subject": "О порядке дежурства",
                "purpose_text": "Установить порядок дежурства.",
                "instruction_items": ["Назначить дежурного.", "Контролировать исполнение."],
                "control_assignee_text": "дежурного офицера",
            },
        },
    )
    assert response.status_code == 201
    assert response.json()["type"] == "INSTRUCTION"
    assert response.json()["acknowledgements"] == []


async def test_create_incoming_and_resolution_and_audit(client, users_fixture, db_session):
    await do_login(client, "incoming_op", "incoming12345")
    incoming_response = await client.post(
        "/api/v1/incoming",
        json={
            "title": "Письмо из штаба",
            "department_id": users_fixture["department"].id,
            "document_date": "2026-05-28",
            "organization_name": "Штаб округа",
            "signer_position": "Командир",
            "signer_name": "И. Иванов",
            "structured_data": {
                "sender": "Штаб округа",
                "received_at": "2026-05-28",
                "subject": "О предоставлении сведений",
                "body_text": "Требуется подготовить сводку.",
            },
        },
    )
    assert incoming_response.status_code == 201
    incoming_id = incoming_response.json()["id"]

    client.cookies.clear()
    await do_login(client, "chief", "chief12345")
    resolution_response = await client.post(
        f"/api/v1/incoming/{incoming_id}/resolution",
        json={
            "title": "Резолюция по письму",
            "department_id": users_fixture["department"].id,
            "document_date": "2026-05-28",
            "organization_name": "ФГКОУ «Екатеринбургское суворовское военное училище»",
            "signer_position": "Начальник училища",
            "signer_name": "А. Кукарцев",
            "structured_data": {
                "linked_incoming_letter_id": incoming_id,
                "resolution_text": "Подготовить ответ до 01.06.2026.",
                "assigned_users": [users_fixture["employee"].id],
                "assigned_departments": [users_fixture["department"].id],
                "assignee_statuses": {},
            },
        },
    )
    assert resolution_response.status_code == 201

    audit_result = await db_session.execute(select(AuditLog).where(AuditLog.action.in_(["incoming.created", "resolution.created"])))
    actions = {item.action for item in audit_result.scalars().all()}
    assert "incoming.created" in actions
    assert "resolution.created" in actions
